import io
import os
import re
import uuid
import logging
from datetime import datetime
from itertools import groupby
from operator import itemgetter
from typing import AsyncGenerator

import numpy as np
import faiss
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer
from openai import OpenAI

try:
    import pytesseract
    from PIL import Image as PilImage, ImageEnhance as PilImageEnhance
    # On Windows the installer drops the binary here but doesn't always add it to PATH
    _WIN_TESS = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    if os.path.exists(_WIN_TESS):
        pytesseract.pytesseract.tesseract_cmd = _WIN_TESS
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

load_dotenv()

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(message)s")
logger = logging.getLogger(__name__)

app = FastAPI(title="RAG Chatbot")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory session store: session_id -> {chunks: [...], index: faiss.Index}
sessions: dict[str, dict] = {}

# Load embedding model once at startup
embedding_model = SentenceTransformer("all-MiniLM-L6-v2")


# ── Helpers ──────────────────────────────────────────────────────────────────

def parse_pdf(data: bytes) -> str:
    doc = fitz.open(stream=data, filetype="pdf")
    pages_text = []
    for page in doc:
        structured_parts = []
        try:
            tabs = page.find_tables()
            for table in tabs.tables:
                rows = []
                for row in table.extract():
                    cells = [str(c).strip() if c is not None else "" for c in row]
                    if any(cells):  # skip completely blank rows
                        rows.append(" | ".join(cells))  # keep empty cells to preserve column position
                if rows:
                    structured_parts.append("[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")
        except AttributeError:
            pass  # PyMuPDF < 1.23 — find_tables not available

        # Use structured output when tables are found; plain text is not used alongside
        # it because the column-mixed plain text is what causes values like OTHERS to
        # bleed into the wrong section (e.g. Deductions).
        if structured_parts:
            pages_text.append("\n\n".join(structured_parts))
        else:
            pages_text.append(page.get_text())
    return "\n\n".join(pages_text)


def parse_pdf_ocr(data: bytes) -> str:
    doc = fitz.open(stream=data, filetype="pdf")
    pages = []

    for page in doc:
        # 300 DPI — industry standard for document OCR; much better than 144 DPI for table text
        pix = page.get_pixmap(matrix=fitz.Matrix(300 / 72, 300 / 72), colorspace=fitz.csRGB)

        # Grayscale + contrast boost: cleans up scan artefacts and faint grid lines
        img = PilImage.open(io.BytesIO(pix.tobytes("png"))).convert("L")
        img = PilImageEnhance.Contrast(img).enhance(1.5)

        # PSM 6 = assume a single uniform block of text (best for mixed table/prose layouts)
        ocr_data = pytesseract.image_to_data(
            img,
            lang="eng",
            config="--oem 3 --psm 6",
            output_type=pytesseract.Output.DICT,
        )

        # Collect words with a confidence floor to discard recognition noise
        words = [
            {
                "text":  ocr_data["text"][i].strip(),
                "x":     ocr_data["left"][i],
                "block": ocr_data["block_num"][i],
                "par":   ocr_data["par_num"][i],
                "line":  ocr_data["line_num"][i],
            }
            for i in range(len(ocr_data["text"]))
            if ocr_data["text"][i].strip() and int(ocr_data["conf"][i]) > 20
        ]

        if not words:
            continue

        # Reconstruct text preserving Tesseract's own row grouping.
        # Within each line the words are sorted left-to-right so "BASIC  15000"
        # stays together rather than being split across phantom columns.
        reconstructed: list[str] = []
        prev_block = None
        line_key = itemgetter("block", "par", "line")
        for (block, _par, _line), group in groupby(sorted(words, key=line_key), key=line_key):
            if prev_block is not None and block != prev_block:
                reconstructed.append("")          # blank line between text blocks
            prev_block = block
            row_text = " ".join(w["text"] for w in sorted(group, key=lambda w: w["x"]))
            reconstructed.append(row_text)

        page_text = re.sub(r"\n{3,}", "\n\n", "\n".join(reconstructed)).strip()
        if page_text:
            pages.append(page_text)

    doc.close()
    return "\n\n".join(pages)


def parse_docx(data: bytes) -> str:
    docx = DocxDocument(io.BytesIO(data))
    return "\n".join(p.text for p in docx.paragraphs if p.text.strip())


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(" ".join(words[start:end]))
        if end == len(words):
            break
        start += chunk_size - overlap
    return chunks


def build_faiss_index(embeddings: np.ndarray) -> faiss.IndexFlatIP:
    # IndexFlatIP = inner product; on L2-normalised vectors this equals cosine similarity
    vectors = embeddings.astype(np.float32)
    faiss.normalize_L2(vectors)
    index = faiss.IndexFlatIP(vectors.shape[1])
    index.add(vectors)
    return index


# ── Endpoints ─────────────────────────────────────────────────────────────────

@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/upload")
async def upload(file: UploadFile = File(...), use_ocr: bool = Form(False)):
    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext not in ("pdf", "docx"):
        raise HTTPException(status_code=400, detail="Unsupported file type. Upload a PDF or DOCX.")

    data = await file.read()
    logger.info("[%s] Upload: %s (ocr=%s)", datetime.now().isoformat(), filename, use_ocr)

    ocr_used = False
    if ext == "pdf" and use_ocr:
        if not OCR_AVAILABLE:
            raise HTTPException(
                status_code=500,
                detail="pytesseract/Pillow not installed. Run: pip install pytesseract Pillow",
            )
        try:
            text = parse_pdf_ocr(data)
            ocr_used = True
        except pytesseract.TesseractNotFoundError:
            raise HTTPException(
                status_code=500,
                detail="Tesseract OCR binary not found. Download the installer from the UB-Mannheim Tesseract releases (tesseract-ocr-w64-setup-*.exe) and keep the default install path (C:\\Program Files\\Tesseract-OCR\\), then restart the server.",
            )
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"OCR failed: {e}")
    else:
        try:
            text = parse_pdf(data) if ext == "pdf" else parse_docx(data)
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"Failed to parse document: {e}")

    if not text.strip():
        raise HTTPException(status_code=422, detail="Document appears to be empty or no text could be extracted.")

    chunks = chunk_text(text)
    embeddings = embedding_model.encode(chunks, show_progress_bar=False, convert_to_numpy=True)

    session_id = str(uuid.uuid4())
    sessions[session_id] = {
        "chunks": chunks,
        "index": build_faiss_index(np.array(embeddings)),
        "filename": filename,
    }

    logger.info("[%s] Indexed %d chunks for session %s (ocr=%s)", datetime.now().isoformat(), len(chunks), session_id, ocr_used)
    return {"session_id": session_id, "chunk_count": len(chunks), "filename": filename, "ocr_used": ocr_used}


class ChatRequest(BaseModel):
    session_id: str
    query: str
    temperature: float = 0.3


# ~100k tokens worth of words — safe limit for gpt-4o-mini's 128k context window
FULL_DOC_WORD_LIMIT = 75_000


async def stream_chat(session: dict, query: str, temperature: float) -> AsyncGenerator[str, None]:
    chunks = session["chunks"]
    total_words = sum(len(c.split()) for c in chunks)

    if total_words <= FULL_DOC_WORD_LIMIT:
        # Document fits in context — pass everything so the LLM can answer any question
        context_chunks = chunks
    else:
        # Large document: use FAISS to retrieve the most relevant 15 chunks
        query_vec = embedding_model.encode([query], show_progress_bar=False, convert_to_numpy=True).astype(np.float32)
        faiss.normalize_L2(query_vec)
        _, top_indices = session["index"].search(query_vec, k=min(15, len(chunks)))
        context_chunks = [chunks[i] for i in top_indices[0]]

    context = "\n\n---\n\n".join(context_chunks)

    system_prompt = (
        "You are a document assistant. Answer the user's question using ONLY the context "
        "excerpts provided below. Do NOT use any outside knowledge.\n"
        "The context may include [TABLE]...[/TABLE] sections where columns are separated by '|'. "
        "Each '|'-delimited cell belongs exclusively to its own column — "
        "NEVER reassign a value from one column (e.g. Earnings or its sub-amounts) to a different column (e.g. Deductions). "
        "For example, if a row reads 'OTHERS | 4878.85 | 4643 | |', the values 4878.85 and 4643 are Earnings figures, "
        "not Deductions, because the Deductions cells in that row are empty.\n"
        "If the answer cannot be found in the provided context, respond with exactly: "
        "\"This information is not available in the uploaded document.\"\n\n"
        f"CONTEXT:\n{context}"
    )

    client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

    stream = client.chat.completions.create(
        model="gpt-4o-mini",
        max_tokens=1024,
        temperature=temperature,
        stream=True,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": query},
        ],
    )

    for chunk in stream:
        delta = chunk.choices[0].delta.content
        if delta:
            # Escape newlines so they survive SSE line-splitting on the client
            yield f"data: {delta.replace(chr(10), '\\n')}\n\n"

    yield "data: [DONE]\n\n"


@app.post("/chat")
async def chat(req: ChatRequest):
    if req.session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found. Please upload a document first.")

    if not req.query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty.")

    temperature = max(0.0, min(1.0, req.temperature))
    logger.info("[%s] Query (session %s): %s", datetime.now().isoformat(), req.session_id, req.query[:80])

    session = sessions[req.session_id]
    return StreamingResponse(
        stream_chat(session, req.query, temperature),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.delete("/session/{session_id}")
def delete_session(session_id: str):
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found.")
    del sessions[session_id]
    logger.info("[%s] Session deleted: %s", datetime.now().isoformat(), session_id)
    return {"status": "deleted"}


# Serve frontend
app.mount("/", StaticFiles(directory="static", html=True), name="static")
